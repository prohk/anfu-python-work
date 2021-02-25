import docx
import os
import re
from docx.shared import Pt
import time
import shutil

# config
file_ls = []
file_dir = input("请输入docx文件的路径")
os.chdir(file_dir)
now_date = time.strftime('%Y-%m-%d')
cnnow_date = now_date.replace("-","年",1).replace("-","月",1)
cnnow_date = cnnow_date + "日"
print("当前日期为:%s"%now_date)
try:
    os.mkdir("result")
    print("文件将输出在result文件夹")
except:
    print("文件已创建")
    shutil.rmtree('result', ignore_errors=True)
    os.mkdir("result")
#文件已创建时会删掉文件及里面文件重新创建
# 正则

re_date = "[\s\S]*(\d{1,4}[年-]\d{1,2}[月-]\d{1,2})[\s\S]*"
# 图案可删
print("*****************批量日期处理 ********************\n\n")
print("<------------记得关掉要处理的文档，否则会读不了出错-------------->\n")
time.sleep(1)
# def
def file_name(file_dir):
    global file_ls
    for root, dirs, files in os.walk(file_dir):
        if files:
            for i in files:
                file_ls.append(i)

def do_table(tables):
    for table in tables:
        rows_num = len(table.rows)
        columns_num = len(table.columns)
        if rows_num == 2 and columns_num == 5:
            cell1 = table.cell(1,1)
            table.cell(1,1).text = "%s"%cnnow_date
            # cell1.paragraphs[0].add_run("")
            # cell1.font.size = 152400


def get_docx(docx_name):
    wd = docx.Document(docx_name)
    para = wd.paragraphs
    tables = wd.tables
    do_table(tables)
    # styles = wd.styles
    # for style in styles:
    #     print(style.name)
    # print(len(para))
    # print(para[32].text)
    for i in para:
        p = wd.paragraphs[17].clear()
        run1 = p.add_run("%s"%cnnow_date)
        font = run1.font
        font.size = Pt(16)
        font.bold =True
        p = wd.paragraphs[49].clear()
        run2 = p.add_run("%s至%s，进行渗透测试；"%(now_date,now_date))
        font = run2.font
        font.size = 152400
        p = wd.paragraphs[50].clear()
        run3 = p.add_run("%s进行报告撰写。"%cnnow_date)
        font = run3.font
        font.size = 152400
        # cmp_result = re.match(re_date,i.text)
        # if cmp_result:
        #     print(i.style.f)

    wd.save("result/%s"%docx_name)


def re_handle(cmpstr, str):
    result = re.match(cmpstr, str)
    if result:
        result = result.group(1)
        return result
    else:
        print("正则出错")
        return


# 获得所有文件名
file_name(file_dir)
print(file_ls)
# 替换



# 读取所有docx
for i in file_ls:
    get_docx(i)